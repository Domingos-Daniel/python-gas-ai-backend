"""
Document processing utilities for handling uploaded files.
Supports Excel, PDF, TXT, and Word documents.
"""
import os
import io
import logging
from typing import Dict, Any, Optional
import pandas as pd
from openpyxl import load_workbook
import PyPDF2
from docx import Document
import mimetypes

# Configure logging
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Processor for various document types with content extraction."""
    
    def __init__(self):
        self.supported_extensions = ['.xlsx', '.xls', '.pdf', '.txt', '.docx', '.doc']
        self.max_file_size = 10 * 1024 * 1024  # 10MB
    
    def process_document(self, file_path: str, file_content: bytes) -> Dict[str, Any]:
        """
        Process a document and extract its content.
        
        Args:
            file_path: Path to the file
            file_content: File content as bytes
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        try:
            # Determine file type
            file_extension = os.path.splitext(file_path)[1].lower()
            mime_type, _ = mimetypes.guess_type(file_path)
            
            if file_extension not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Check file size
            if len(file_content) > self.max_file_size:
                raise ValueError(f"File too large. Maximum size: {self.max_file_size / (1024*1024):.1f}MB")
            
            # Process based on file type
            if file_extension in ['.xlsx', '.xls']:
                return self._process_excel(file_content, file_path)
            elif file_extension == '.pdf':
                return self._process_pdf(file_content, file_path)
            elif file_extension == '.txt':
                return self._process_txt(file_content, file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._process_word(file_content, file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise
    
    def _process_excel(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process Excel files and extract data."""
        try:
            # Create a BytesIO object from the content
            excel_file = io.BytesIO(content)
            
            # Read all sheets
            xl_file = pd.ExcelFile(excel_file)
            sheets_data = {}
            
            for sheet_name in xl_file.sheet_names:
                try:
                    # Read the sheet
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    
                    # Convert to dict, handling NaN values
                    sheet_data = df.replace({pd.NA: None, pd.NaT: None, float('nan'): None})
                    sheets_data[sheet_name] = sheet_data.to_dict('records')
                    
                except Exception as e:
                    logger.warning(f"Error reading sheet '{sheet_name}': {str(e)}")
                    sheets_data[sheet_name] = []
            
            # Extract text content from first sheet for context
            text_content = ""
            if xl_file.sheet_names:
                first_sheet = pd.read_excel(excel_file, sheet_name=xl_file.sheet_names[0])
                text_content = f"Excel file '{filename}' contains {len(xl_file.sheet_names)} sheets. "
                text_content += f"First sheet '{xl_file.sheet_names[0]}' has {len(first_sheet)} rows and {len(first_sheet.columns)} columns. "
                text_content += f"Columns: {', '.join(first_sheet.columns.astype(str).tolist()[:10])}"
                if len(first_sheet.columns) > 10:
                    text_content += "..."
            
            return {
                'type': 'excel',
                'filename': filename,
                'sheets': sheets_data,
                'sheet_names': xl_file.sheet_names,
                'text_content': text_content,
                'metadata': {
                    'total_sheets': len(xl_file.sheet_names),
                    'file_size': len(content)
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing Excel file {filename}: {str(e)}")
            raise ValueError(f"Error processing Excel file: {str(e)}")
    
    def _process_pdf(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF files and extract text."""
        try:
            # Create a BytesIO object from the content
            pdf_file = io.BytesIO(content)
            
            # Create PDF reader object
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract text from all pages
            text_content = ""
            page_count = len(pdf_reader.pages)
            
            for page_num in range(page_count):
                try:
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
            
            # Clean up the text
            text_content = text_content.strip()
            
            # Extract metadata if available
            metadata = {}
            if pdf_reader.metadata:
                metadata = {
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'subject': pdf_reader.metadata.get('/Subject', ''),
                    'creator': pdf_reader.metadata.get('/Creator', ''),
                    'producer': pdf_reader.metadata.get('/Producer', ''),
                    'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                    'modification_date': str(pdf_reader.metadata.get('/ModDate', ''))
                }
            
            return {
                'type': 'pdf',
                'filename': filename,
                'text_content': text_content,
                'metadata': {
                    'page_count': page_count,
                    'file_size': len(content),
                    **metadata
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing PDF file {filename}: {str(e)}")
            raise ValueError(f"Error processing PDF file: {str(e)}")
    
    def _process_txt(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process text files."""
        try:
            # Decode the content
            text_content = content.decode('utf-8')
            
            # Count lines and characters
            lines = text_content.split('\n')
            
            return {
                'type': 'txt',
                'filename': filename,
                'text_content': text_content,
                'metadata': {
                    'line_count': len(lines),
                    'char_count': len(text_content),
                    'file_size': len(content)
                }
            }
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                text_content = content.decode('latin-1')
                lines = text_content.split('\n')
                
                return {
                    'type': 'txt',
                    'filename': filename,
                    'text_content': text_content,
                    'metadata': {
                        'line_count': len(lines),
                        'char_count': len(text_content),
                        'file_size': len(content),
                        'encoding': 'latin-1'
                    }
                }
            except Exception as e:
                logger.error(f"Error decoding text file {filename}: {str(e)}")
                raise ValueError(f"Error decoding text file: {str(e)}")
        except Exception as e:
            logger.error(f"Error processing text file {filename}: {str(e)}")
            raise ValueError(f"Error processing text file: {str(e)}")
    
    def _process_word(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process Word documents."""
        try:
            # Create a BytesIO object from the content
            word_file = io.BytesIO(content)
            
            # Load the document
            doc = Document(word_file)
            
            # Extract text from all paragraphs
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            table_text = ""
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    if any(row_text):
                        table_text += " | ".join(row_text) + "\n"
            
            # Combine paragraph and table text
            full_text = text_content
            if table_text:
                full_text += "\n\nTabelas:\n" + table_text
            
            # Extract basic metadata
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'file_size': len(content)
            }
            
            return {
                'type': 'word',
                'filename': filename,
                'text_content': full_text.strip(),
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing Word file {filename}: {str(e)}")
            raise ValueError(f"Error processing Word file: {str(e)}")
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic information about a file."""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            mime_type, _ = mimetypes.guess_type(file_path)
            
            return {
                'extension': file_extension,
                'mime_type': mime_type,
                'supported': file_extension in self.supported_extensions
            }
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {str(e)}")
            return {
                'extension': '',
                'mime_type': None,
                'supported': False,
                'error': str(e)
            }

# Create a global instance
document_processor = DocumentProcessor()

def process_uploaded_document(file_path: str, file_content: bytes) -> Dict[str, Any]:
    """
    Convenience function to process an uploaded document.
    
    Args:
        file_path: Path to the uploaded file
        file_content: File content as bytes
        
    Returns:
        Dictionary containing extracted content and metadata
    """
    return document_processor.process_document(file_path, file_content)